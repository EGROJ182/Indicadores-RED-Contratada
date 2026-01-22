import pandas as pd
from docx import Document
import mysql.connector
from datetime import date, datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
import shutil
import os
import sys
import time

# --- Configuración de rutas ---
BASE_DIR = r'D:\Proyectos\Indicadores\Resources'
PATH_LOGS = r'D:\Proyectos\Indicadores\Logs'
BASE_DIR_OUTPUT = r'D:\Proyectos\Indicadores\Salidas'
RUTA_LOGS = PATH_LOGS
RUTA_IMAGEN = os.path.join(BASE_DIR, 'f.png')
RUTA_READING = os.path.join(BASE_DIR, 'reading.xlsx')
RUTA_INDICADORES = os.path.join(BASE_DIR, 'indicadores_anexo_9.xlsx')
RUTA_BORRADOR = os.path.join(BASE_DIR, 'borrador_indicadores.docx')
# --- Ruta One drive para almacenar copia de los Anexos 9 generados ---
COPIE_DIR = os.path.join(r'C:\Users\JORGEEDILSONVEGAACOS\One Drive Analista Red 11\OneDrive - Positiva Compañia de Seguros S. A\Indicadores\Anexos9')

# Función para obtener la ruta de la base maestra o nuevos proveedores del año en curso
def obtener_ruta_archivo():
    while True:
        try:
            print("\nSelecciona el archivo a procesar:")
            print("1. proveedores.xlsx")
            print("2. nuevos.xlsx")
            print("3. Salir")
            
            opcion = input("Ingresa tu opción (1 o 2): ").strip()
            
            if opcion == "1":
                return os.path.join(BASE_DIR, 'proveedores.xlsx')
            elif opcion == "2":
                return os.path.join(BASE_DIR, 'nuevos.xlsx')
            elif opcion == "3":
                print("\n\n⚠️ Operación cancelada por el usuario.")
                exit()
            else:
                raise ValueError("Opción no válida. Debes elegir 1 o 2.")
                
        except ValueError as e:
            print(f"❌ Error: {e}")
            print("Por favor, intenta nuevamente.\n")
        except KeyboardInterrupt:
            print("\n\n⚠️ Operación cancelada por el usuario.")
            exit()

# Usar la función
RUTA_PROVEEDORES = obtener_ruta_archivo()

# --- Configuración de la Base de Datos ---
db_config = {
    'host': 'localhost',
    'port': 3306,
    'user': 'root',
    'password': '!1q2w3e4r5t6y7u8i9o0p',
    'database': 'data_tarifas_reps'
}

# --- Variables globales para logs ---
logs_procesamiento = []

def crear_nombre_archivo_log():
    """
    Crea el nombre del archivo de log con formato log_diamesañohoraminutosegundo
    """
    ahora = datetime.now()
    nombre = f"log_{ahora.day:02d}{ahora.month:02d}{ahora.year}{ahora.hour:02d}{ahora.minute:02d}{ahora.second:02d}.txt"
    return nombre

def agregar_log(nit, estado):
    """
    Agrega un registro al log de procesamiento
    """
    logs_procesamiento.append({
        'nit': nit,
        'estado': estado,
        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    })

def guardar_logs():
    """
    Guarda los logs en un archivo de texto en la carpeta LOGS
    """
    try:
        # Verificar que PATH_LOGS existe o crearlo
        if not os.path.exists(PATH_LOGS):
            os.makedirs(PATH_LOGS, exist_ok=True)
        # os.makedirs(RUTA_LOGS, exist_ok=True)
        nombre_archivo = crear_nombre_archivo_log()
        ruta_log = os.path.join(RUTA_LOGS, nombre_archivo)
        
        with open(ruta_log, 'w', encoding='utf-8') as f:
            f.write(f"LOG DE PROCESAMIENTO - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("=" * 70 + "\n\n")
            
            for log in logs_procesamiento:
                f.write(f"NIT: {log['nit']} - {log['estado']} - {log['timestamp']}\n")
        
        return ruta_log
    except Exception as e:
        print(f"\nError al guardar logs: {e}")
        return None

def imprimir_dinamico(texto, linea_nueva=False):
    """
    Imprime texto de forma dinámica, actualizando la misma línea
    """
    if linea_nueva:
        print(f"\n{texto}")
    else:
        print(f"\r{texto}", end='', flush=True)

def limpiar_linea():
    """
    Limpia la línea actual de la consola
    """
    print('\r' + ' ' * 100 + '\r', end='', flush=True)

def formatear_fecha(valor, fmt="%Y-%m-%d"):
    """
    Devuelve la fecha en formato yyyy/mm/dd.
    Acepta pandas.Timestamp, datetime/date, string o seriales de Excel.
    Si no puede convertir, devuelve str(valor).
    """
    if valor is None or (isinstance(valor, float) and pd.isna(valor)):
        return ""
    if isinstance(valor, (pd.Timestamp, datetime, date)):
        return pd.to_datetime(valor).strftime(fmt)
    dt = pd.to_datetime(valor, errors="coerce")
    if pd.isna(dt):
        return str(valor) if pd.notna(valor) else ""
    return dt.strftime(fmt)

# --- Funciones Auxiliares ---
def obtener_codigos_reps_habilitados(nit):
    """
    Consulta la base de datos para obtener los códigos REPS habilitados para un NIT.
    """
    codigos_habilitados = set()
    try:
        conn = mysql.connector.connect(**db_config)
        cursor = conn.cursor()
        query = f"SELECT DISTINCT cr.serv_codigo FROM consolidado_reps cr WHERE cr.nit = '{nit}'"
        cursor.execute(query)
        for (serv_codigo,) in cursor:
            codigos_habilitados.add(str(serv_codigo).strip())
        cursor.close()
        conn.close()
    except mysql.connector.Error as err:
        raise Exception(f"Error al conectar a la base de datos: {err}")
    return codigos_habilitados

def es_numerico(s):
    """
    Verifica si una cadena de texto es completamente numérica.
    """
    return str(s).isdigit()

def consolidar_indicadores_duplicados(indicadores_aplicables_df, codigos_reps_habilitados):
    """
    Consolida indicadores que tienen el mismo nombre_indicador concatenando cod_propio y cod_reps.
    Retorna una lista de diccionarios con los indicadores consolidados.
    """
    indicadores_consolidados = {}
    
    for ind_index, ind_row in indicadores_aplicables_df.iterrows():
        cod_reps = str(ind_row['cod_reps']).strip()
        
        # Validar si el indicador debe aplicarse
        aplicar_indicador = False
        if es_numerico(cod_reps):
            if cod_reps in codigos_reps_habilitados:
                aplicar_indicador = True
        else:
            aplicar_indicador = True
        
        if aplicar_indicador:
            nombre_indicador = str(ind_row['nombre_indicador']).strip()
            cod_propio = str(ind_row['cod_propio']).strip()
            
            # Si ya existe un indicador con este nombre
            if nombre_indicador in indicadores_consolidados:
                # Concatenar cod_propio
                cod_propio_existente = indicadores_consolidados[nombre_indicador]['cod_propio']
                if cod_propio not in cod_propio_existente:
                    indicadores_consolidados[nombre_indicador]['cod_propio'] = f"{cod_propio_existente}, {cod_propio}"
                
                # Concatenar cod_reps
                cod_reps_existente = indicadores_consolidados[nombre_indicador]['cod_reps']
                if cod_reps not in cod_reps_existente:
                    indicadores_consolidados[nombre_indicador]['cod_reps'] = f"{cod_reps_existente}, {cod_reps}"
            else:
                # Crear nuevo indicador consolidado
                indicador_consolidado = dict(ind_row)
                indicador_consolidado['cod_propio'] = cod_propio
                indicador_consolidado['cod_reps'] = cod_reps
                indicadores_consolidados[nombre_indicador] = indicador_consolidado
    
    # Convertir el diccionario a lista para mantener compatibilidad con el resto del código
    return list(indicadores_consolidados.values())

def agregar_imagen_ficha_tecnica(documento):
    """
    Agrega la imagen f.png como encabezado de ficha técnica.
    """
    try:
        if os.path.exists(RUTA_IMAGEN):
            # Crear un nuevo párrafo para la imagen
            paragraph = documento.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Eliminar espaciado antes y después del párrafo
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            
            # Agregar la imagen al párrafo
            run = paragraph.add_run()
            # Ancho igual al de la tabla (4.8 + 11.2 = 16 cm) y alto de 0.7 cm
            run.add_picture(RUTA_IMAGEN, width=Cm(16), height=Cm(0.7))
            
            return True
        else:
            raise Exception(f"No se encontró el archivo f.png en la ruta: {RUTA_IMAGEN}")
            
    except Exception as e:
        raise Exception(f"Error al agregar imagen f.png: {e}")

def agregar_bordes_tabla(tabla):
    """
    Agrega solo bordes horizontales internos a la tabla (entre filas).
    """
    try:
        # Obtener el elemento XML de la tabla
        tbl = tabla._element
        
        # Crear elemento de bordes
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Crear bordes de tabla
        tblBorders = OxmlElement('w:tblBorders')
        
        # Definir estilos de borde
        border_style = "single"
        border_size = "4"  # Tamaño del borde
        border_color = "000000"  # Color negro
        
        # Solo crear borde horizontal interno (entre filas)
        insideH = OxmlElement('w:insideH')
        insideH.set(qn('w:val'), border_style)
        insideH.set(qn('w:sz'), border_size)
        insideH.set(qn('w:color'), border_color)
        tblBorders.append(insideH)
        
        # Eliminar bordes existentes si los hay
        existing_borders = tblPr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tblPr.remove(existing_borders)
            
        tblPr.append(tblBorders)
        
    except Exception as e:
        raise Exception(f"No se pudieron agregar bordes a la tabla: {e}")

def copiar_formato_tabla_simple(tabla_origen, tabla_destino):
    """
    Copia el formato básico de una tabla origen a una tabla destino con solo 2 columnas.
    """
    try:
        # Copiar el estilo de la tabla
        if tabla_origen.style:
            tabla_destino.style = tabla_origen.style
        
        # Configurar el ancho total de la tabla
        tabla_destino.autofit = False
        
        # Ajustar el ancho de las columnas: 30% para títulos, 70% para contenido
        if len(tabla_destino.columns) == 2:
            try:
                from docx.shared import Cm
                # Usar centímetros para mayor precisión
                # Ancho total aproximado de 16cm, 30% = 4.8cm, 70% = 11.2cm
                tabla_destino.columns[0].width = Cm(4.8)   # 30%
                tabla_destino.columns[1].width = Cm(11.2)  # 70%
                
                # Forzar el ancho a nivel XML para mayor efectividad
                tbl = tabla_destino._element
                tblGrid = tbl.find(qn('w:tblGrid'))
                if tblGrid is not None:
                    gridCols = tblGrid.findall(qn('w:gridCol'))
                    if len(gridCols) >= 2:
                        gridCols[0].set(qn('w:w'), '2304')  # 30% en twips (4.8cm * 480)
                        gridCols[1].set(qn('w:w'), '5376')  # 70% en twips (11.2cm * 480)
                        
            except Exception as e:
                raise Exception(f"No se pudo ajustar el ancho de las columnas: {e}")
                
    except Exception as e:
        raise Exception(f"No se pudo copiar completamente el formato de la tabla: {e}")

def copiar_formato_tabla(tabla_origen, tabla_destino):
    """
    Copia el formato de una tabla origen a una tabla destino.
    """
    try:
        # Copiar el estilo de la tabla
        if tabla_origen.style:
            tabla_destino.style = tabla_origen.style
        
        # Copiar propiedades de celdas
        for r_idx, row_origen in enumerate(tabla_origen.rows):
            if r_idx < len(tabla_destino.rows):
                for c_idx, cell_origen in enumerate(row_origen.cells):
                    if c_idx < len(tabla_destino.rows[r_idx].cells):
                        cell_destino = tabla_destino.rows[r_idx].cells[c_idx]
                        # Copiar el ancho de la columna si es posible
                        try:
                            cell_destino.width = cell_origen.width
                        except:
                            pass
    except Exception as e:
        raise Exception(f"No se pudo copiar completamente el formato de la tabla: {e}")

def reemplazar_placeholders_en_tabla(tabla, indicador, supervisor, ccm):
    """
    Reemplaza los placeholders en una tabla con los datos del indicador.
    """
    # Diccionario de reemplazos
    reemplazos = {
        '[tipo_proveedor]': str(ccm),
        '[tipo_indicador]': str(indicador.get('tipo_indicador', '')),
        '[categoria]': str(indicador.get('categoria', '')),
        '[nombre_indicador]': str(indicador.get('nombre_indicador', '')),
        '[fecha_creacion_indicador]': str(indicador.get('fecha_creacion_indicador', '')),
        '[cod_res_256_de_2016]': str(indicador.get('cod_res_256_de_2016', '')),
        '[cod_propio]': str(indicador.get('cod_propio', '')),
        '[descripcion]': str(indicador.get('descripcion', '')),
        '[formula]': str(indicador.get('formula', '')),
        '[numerador]': str(indicador.get('numerador', '')),
        '[fuente_numerador]': str(indicador.get('fuente_numerador', '')),
        '[denominador]': str(indicador.get('denominador', '')),
        '[fuente_denominador]': str(indicador.get('fuente_denominador', '')),
        '[unidad_de_medida]': str(indicador.get('unidad_de_medida', '')),
        '[meta]': str(indicador.get('meta', '')),
        '[periodicidad]': str(indicador.get('periodicidad', '')),
        '[progresividad]': str(indicador.get('progresividad', '')),
        '[observaciones]': str(indicador.get('observaciones', '')),
        '[metodologia]': str(indicador.get('metodologia', '')),
        '[exclusiones]': str(indicador.get('exclusiones', '')),
        '[responsable]': str(indicador.get('responsable', '')),
        '[cod_reps]': str(indicador.get('cod_reps', '')),
        '[grupo_indicador]': str(indicador.get('grupo_indicador', ''))
    }
    
    # Reemplazar en todas las celdas
    for row in tabla.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                texto_original = paragraph.text
                for placeholder, valor in reemplazos.items():
                    texto_original = texto_original.replace(placeholder, valor)
                paragraph.text = texto_original

def llenar_tabla_proveedor(tabla, proveedor_data):
    """
    Llena la tabla de información del proveedor.
    """
    # --- FECHAS del proveedor formateadas ---
    fin_vig = proveedor_data.get('fin_vigencia', '')

    # Mapeo de placeholders a valores del proveedor
    reemplazos_proveedor = {
        '[sucursal]': str(proveedor_data.get('sucursal', '')),
        '[nombre]': str(proveedor_data.get('nombre', '')),
        '[nit]': str(proveedor_data.get('nit', '')),
        '[numero_contrato]': str(proveedor_data.get('numero_contrato', '')),
        '[year_contrato]': str(proveedor_data.get('year_contrato', '')),
        '[fin_vigencia]': formatear_fecha(fin_vig),
        '[categoría_cuentas_medicas]': str(proveedor_data.get('categoria_cuentas_medicas', '')),
        '[supervisor]': 'Supervisor y/o interventoría designada'
    }
    
    # Reemplazar en todas las celdas de la tabla
    for row in tabla.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                texto_original = paragraph.text
                for placeholder, valor in reemplazos_proveedor.items():
                    texto_original = texto_original.replace(placeholder, valor)
                paragraph.text = texto_original

def agregar_salto_de_pagina(documento):
    """
    Agrega un salto de página al documento.
    """
    paragraph = documento.add_paragraph()
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    fldChar = OxmlElement('w:br')
    fldChar.set(qn('w:type'), 'page')
    run._r.append(fldChar)

def extraer_contenido_cierre_y_firmas(documento_borrador):
    """
    Extrae el párrafo de cierre y las firmas del documento borrador.
    Retorna una tupla: (parrafo_cierre, lista_de_paragrafos_firmas)
    """
    parrafo_cierre = None
    parrafos_firmas = []
    encontro_cierre = False
    
    for paragraph in documento_borrador.paragraphs:
        texto = paragraph.text.strip()
        
        # Buscar el párrafo de cierre
        if "Positiva compañía de seguros, mediante la firma de interventoría" in texto:
            parrafo_cierre = paragraph  # Guardar el párrafo completo, no solo el texto
            encontro_cierre = True
            continue
        
        # Si ya encontró el cierre, los siguientes párrafos son las firmas
        if encontro_cierre and texto:
            # Guardar el párrafo completo para mantener el formato
            parrafos_firmas.append(paragraph)
    
    return parrafo_cierre, parrafos_firmas

def copiar_parrafo_con_formato(documento_destino, parrafo_origen):
    """
    Copia un párrafo completo con todo su formato al documento destino.
    """
    try:
        # Crear nuevo párrafo en el documento destino
        nuevo_parrafo = documento_destino.add_paragraph()
        
        # Copiar formato del párrafo
        nuevo_parrafo.alignment = parrafo_origen.alignment
        nuevo_parrafo.paragraph_format.space_before = parrafo_origen.paragraph_format.space_before
        nuevo_parrafo.paragraph_format.space_after = parrafo_origen.paragraph_format.space_after
        nuevo_parrafo.paragraph_format.line_spacing = parrafo_origen.paragraph_format.line_spacing
        nuevo_parrafo.paragraph_format.first_line_indent = parrafo_origen.paragraph_format.first_line_indent
        nuevo_parrafo.paragraph_format.left_indent = parrafo_origen.paragraph_format.left_indent
        nuevo_parrafo.paragraph_format.right_indent = parrafo_origen.paragraph_format.right_indent
        
        # Copiar estilo del párrafo si existe
        if parrafo_origen.style:
            try:
                nuevo_parrafo.style = parrafo_origen.style
            except:
                pass
        
        # Copiar cada run (fragmento de texto con formato) del párrafo original
        for run_origen in parrafo_origen.runs:
            nuevo_run = nuevo_parrafo.add_run(run_origen.text)
            
            # Copiar formato del run
            nuevo_run.bold = run_origen.bold
            nuevo_run.italic = run_origen.italic
            nuevo_run.underline = run_origen.underline
            
            # Copiar fuente si existe
            if run_origen.font.name:
                nuevo_run.font.name = run_origen.font.name
            if run_origen.font.size:
                nuevo_run.font.size = run_origen.font.size
            if run_origen.font.color.rgb:
                nuevo_run.font.color.rgb = run_origen.font.color.rgb
        
        return nuevo_parrafo
        
    except Exception as e:
        raise Exception(f"Error copiando formato del párrafo: {e}")

def limpiar_nombre_archivo(nombre):
    """
    Limpia caracteres inválidos de un nombre de archivo y limita su longitud.
    """
    # Caracteres inválidos en Windows
    caracteres_invalidos = r'<>:"/\|?*'
    
    # Reemplazar caracteres inválidos por guion bajo o espacio
    nombre_limpio = nombre
    for char in caracteres_invalidos:
        if char == '/':
            nombre_limpio = nombre_limpio.replace(char, '-')  # Reemplazar / por -
        else:
            nombre_limpio = nombre_limpio.replace(char, '_')
    
    # Eliminar espacios múltiples
    import re
    nombre_limpio = re.sub(r'\s+', ' ', nombre_limpio)
    
    # Eliminar espacios al inicio y final
    nombre_limpio = nombre_limpio.strip()
    
    # Limitar longitud a 150 caracteres para evitar problemas con rutas largas
    if len(nombre_limpio) > 150:
        nombre_limpio = nombre_limpio[:150].strip()
    
    return nombre_limpio

def verificar_archivos_requeridos():
    """
    Verifica que todos los archivos requeridos existan.
    """
    archivos_requeridos = {
        'reading.xlsx': RUTA_READING,
        'proveedores.xlsx': RUTA_PROVEEDORES, 
        'indicadores_anexo_9.xlsx': RUTA_INDICADORES,
        'borrador_indicadores.docx': RUTA_BORRADOR,
        'f.png': RUTA_IMAGEN
    }
    
    archivos_faltantes = []
    for nombre, ruta in archivos_requeridos.items():
        if not os.path.exists(ruta):
            archivos_faltantes.append(f"{nombre} -> {ruta}")
    
    if archivos_faltantes:
        raise Exception(f"No se encontraron los siguientes archivos: {', '.join(archivos_faltantes)}")

def copiar_word_a_onedrive(ruta_word):
    """
    Copia el archivo Word a la carpeta de OneDrive.
    """
    try:
        os.makedirs(COPIE_DIR, exist_ok=True)
        destino = os.path.join(COPIE_DIR, os.path.basename(ruta_word))
        shutil.copy2(ruta_word, destino)
        return True
    except Exception as e:
        raise Exception(f"Error copiando Word a OneDrive: {e}")

# --- Función Principal ---
def generar_documentos_indicadores():
    """
    Función principal que genera los documentos con indicadores.
    """
    print("🚀 GENERADOR DE DOCUMENTOS DE INDICADORES (Solo Word)")
    print("=" * 60)
    print(f"📁 Directorio base: {BASE_DIR}")
    print(f"📁 Directorio salida: {BASE_DIR_OUTPUT}")
    print(f"📁 OneDrive: {COPIE_DIR}")
    
    # 1. Verificar archivos necesarios
    try:
        imprimir_dinamico("🔍 Verificando archivos requeridos...")
        verificar_archivos_requeridos()
        limpiar_linea()
        print("✅ Todos los archivos requeridos están disponibles.")
    except Exception as e:
        limpiar_linea()
        print(f"❌ {e}")
        return
    
    # 2. Crear directorios de salida si no existen
    os.makedirs(BASE_DIR_OUTPUT, exist_ok=True)
    os.makedirs(COPIE_DIR, exist_ok=True)
    os.makedirs(RUTA_LOGS, exist_ok=True)
    
    # 3. Cargar los archivos de Excel
    try:
        imprimir_dinamico("📊 Cargando archivos de Excel...")
        df_borrador = pd.read_excel(RUTA_READING, sheet_name='quering')
        df_proveedores = pd.read_excel(RUTA_PROVEEDORES, sheet_name='proveedores')
        df_indicadores_anexo_9 = pd.ExcelFile(RUTA_INDICADORES)
        df_hom_ccm_ind = df_indicadores_anexo_9.parse('hom_ccm_ind')
        df_indicadores = df_indicadores_anexo_9.parse('indicadores')
        limpiar_linea()
        print("✅ Archivos cargados exitosamente")
    except Exception as e:
        limpiar_linea()
        print(f"❌ Error al cargar archivos: {e}")
        return

    # 4. Cargar el documento borrador y extraer contenido de cierre y firmas
    try:
        imprimir_dinamico("📄 Cargando documento borrador...")
        doc_borrador_plantilla = Document(RUTA_BORRADOR)
        parrafo_cierre, parrafos_firmas = extraer_contenido_cierre_y_firmas(doc_borrador_plantilla)
        limpiar_linea()
        print(f"✅ Documento borrador cargado - {len(parrafos_firmas)} párrafos de firmas extraídos")
    except Exception as e:
        limpiar_linea()
        print(f"❌ Error al cargar el documento borrador: {e}")
        return

    # 5. Iterar sobre cada NIT del borrador
    documentos_generados = 0
    archivos_copiados_onedrive = 0
    
    total_nits = len(df_borrador)
    print(f"\n🔄 Procesando {total_nits} proveedores...")
    print("")  # Línea extra para el contador dinámico
    
    for index, row_borrador in df_borrador.iterrows():
        nit = str(row_borrador['nit']).strip()
        categoria_ccm_borrador = str(row_borrador['categoria_cuentas_medicas']).strip()
        
        # Mostrar progreso dinámico
        progreso = f"[{index + 1}/{total_nits}] 🏢 NIT: {nit}"
        imprimir_dinamico(progreso)

        try:
            # Buscar el NIT en el archivo de proveedores
            proveedor_data = df_proveedores[df_proveedores['nit'].astype(str).str.strip() == nit]
            if proveedor_data.empty:
                agregar_log(nit, "Error - NIT no encontrado en proveedores.xlsx")
                continue

            proveedor_row = proveedor_data.iloc[0]
            supervisor_proveedor = str(proveedor_row['supervisor'])
            ccm = str(proveedor_row['categoria_cuentas_medicas'])

            # Generar el nombre de archivo de salida
            nombre_proveedor = limpiar_nombre_archivo(str(proveedor_row['nombre']))
            num_contrato = str(proveedor_row['numero_contrato'])
            year_contrato = str(proveedor_row['year_contrato'])
            nit_proveedor = str(proveedor_row['nit'])
            nombre_archivo_word = f"Anexo 9 {num_contrato}-{year_contrato} {nombre_proveedor} {nit_proveedor}.docx"
            
            # Crear la ruta completa del archivo de salida
            ruta_archivo_word = os.path.join(BASE_DIR_OUTPUT, nombre_archivo_word)

            # Actualizar progreso con más detalle
            imprimir_dinamico(f"{progreso} - Creando documento...")

            # Crear una copia del borrador original
            shutil.copyfile(RUTA_BORRADOR, ruta_archivo_word)
            documento_final = Document(ruta_archivo_word)

            # Llenar la información del proveedor (primera tabla)
            if len(documento_final.tables) > 0:
                tabla_info_proveedor = documento_final.tables[0]
                llenar_tabla_proveedor(tabla_info_proveedor, proveedor_row)

            # Limpiar párrafos de cierre y firmas del documento original
            paragrafos_a_eliminar = []
            eliminar_resto = False
            
            for paragraph in documento_final.paragraphs:
                texto = paragraph.text.strip()
                
                # Si encuentra el párrafo de cierre, marcar para eliminar este y todos los siguientes
                if "Positiva compañía de seguros, mediante la firma de interventoría" in texto:
                    eliminar_resto = True
                
                # Si ya marcamos para eliminar, agregar todos los párrafos siguientes
                if eliminar_resto:
                    paragrafos_a_eliminar.append(paragraph)

            # Eliminar los párrafos encontrados
            for paragraph in paragrafos_a_eliminar:
                p = paragraph._element
                p.getparent().remove(p)

            # Actualizar progreso
            imprimir_dinamico(f"{progreso} - Consultando indicadores...")

            # Encontrar y validar los indicadores
            grupo_ind_data = df_hom_ccm_ind[df_hom_ccm_ind['categoria_cuentas_medicas'].astype(str).str.strip() == categoria_ccm_borrador]
            if grupo_ind_data.empty:
                agregar_log(nit, f"Error - Categoría '{categoria_ccm_borrador}' no encontrada en indicadores")
                continue

            grupo_indicador = str(grupo_ind_data.iloc[0]['grupo_indicador']).strip()
            
            # Definir grupos a buscar
            grupos_a_buscar = [grupo_indicador, 'TODOS LOS PROVEEDORES']
            if grupo_indicador == 'IPS':
                grupos_a_buscar.extend(['IPS CON RHB', 'IPS-RHI-RIESGO BIOLOGICO'])
            elif grupo_indicador == 'ALIADO RHI':
                grupos_a_buscar.extend(['IPS-RHI-RIESGO BIOLOGICO'])
            elif grupo_indicador == 'RIESGO BIOLOGICO':
                grupos_a_buscar.extend(['IPS-RHI-RIESGO BIOLOGICO'])
            elif grupo_indicador == 'MEDICAMENTOS':
                grupos_a_buscar.extend(['MEDICAMENTOS'])

            indicadores_aplicables_df = df_indicadores[df_indicadores['grupo_indicador'].astype(str).str.strip().isin(grupos_a_buscar)]

            # Actualizar progreso
            imprimir_dinamico(f"{progreso} - Consultando REPS...")
            
            codigos_reps_habilitados = obtener_codigos_reps_habilitados(nit)

            # NUEVA FUNCIONALIDAD: Consolidar indicadores con mismo nombre_indicador
            indicadores_consolidados = consolidar_indicadores_duplicados(indicadores_aplicables_df, codigos_reps_habilitados)
            
            # Obtener la tabla plantilla para las fichas (la última tabla del documento)
            if len(documento_final.tables) > 1:
                ficha_plantilla = documento_final.tables[-1]
                
                # Remover la tabla plantilla original para evitar duplicados
                tabla_elemento = ficha_plantilla._element
                tabla_elemento.getparent().remove(tabla_elemento)
            else:
                agregar_log(nit, "Error - No se encontró tabla plantilla para fichas")
                continue

            # Actualizar progreso
            imprimir_dinamico(f"{progreso} - Generando fichas ({len(indicadores_consolidados)} indicadores)...")

            # Generar fichas de indicadores consolidados
            indicadores_procesados = 0
            
            for indicador in indicadores_consolidados:
                nombre_indicador = str(indicador['nombre_indicador']).strip()
                cod_propio = str(indicador['cod_propio']).strip()
                cod_reps = str(indicador['cod_reps']).strip()
                
                # Agregar la imagen de "ficha técnica" antes de cada tabla
                agregar_imagen_ficha_tecnica(documento_final)
                
                # Crear nueva tabla con solo 2 columnas (etiqueta y valor)
                nueva_tabla = documento_final.add_table(
                    rows=len(ficha_plantilla.rows),
                    cols=2  # Solo 2 columnas: etiqueta y valor
                )
                
                # Copiar formato de la tabla original
                copiar_formato_tabla_simple(ficha_plantilla, nueva_tabla)
                
                # Agregar bordes a la tabla
                agregar_bordes_tabla(nueva_tabla)
                
                # Copiar solo el contenido de las primeras 2 columnas de la plantilla
                filas_copiadas = 0
                for r_idx, row_plantilla in enumerate(ficha_plantilla.rows):
                    # Verificar si la fila tiene contenido útil (no está vacía)
                    tiene_contenido = False
                    for c_idx in range(min(2, len(row_plantilla.cells))):
                        if row_plantilla.cells[c_idx].text.strip():
                            tiene_contenido = True
                            break
                    
                    # Solo copiar filas que tengan contenido
                    if tiene_contenido and filas_copiadas < len(nueva_tabla.rows):
                        for c_idx in range(min(2, len(row_plantilla.cells))):
                            if c_idx < len(nueva_tabla.rows[filas_copiadas].cells):
                                nueva_tabla.cell(filas_copiadas, c_idx).text = row_plantilla.cells[c_idx].text.strip()
                        filas_copiadas += 1
                
                # Si sobran filas en la nueva tabla, eliminarlas
                while len(nueva_tabla.rows) > filas_copiadas:
                    row_to_remove = nueva_tabla.rows[-1]._element
                    row_to_remove.getparent().remove(row_to_remove)

                # Reemplazar placeholders con datos del indicador consolidado
                reemplazar_placeholders_en_tabla(nueva_tabla, indicador, supervisor_proveedor, ccm)
                
                # Agregar espacio entre fichas
                documento_final.add_paragraph()
                
                indicadores_procesados += 1

            # Actualizar progreso
            imprimir_dinamico(f"{progreso} - Agregando firmas...")

            # Agregar el párrafo de cierre y las firmas al final del documento
            if parrafo_cierre:
                # Agregar el párrafo de cierre con su formato original
                copiar_parrafo_con_formato(documento_final, parrafo_cierre)
                
                # Agregar espacio antes de las firmas
                documento_final.add_paragraph()
                
                # Agregar cada párrafo de firma con su formato original completo
                for i, parrafo_firma in enumerate(parrafos_firmas):
                    copiar_parrafo_con_formato(documento_final, parrafo_firma)
                    
                    # Si no es la última firma, agregar espacio entre firmas
                    # Detectar cambio de firma (de POSITIVA a CONTRATISTA)
                    if i < len(parrafos_firmas) - 1:
                        texto_actual = parrafo_firma.text.strip().upper()
                        texto_siguiente = parrafos_firmas[i + 1].text.strip().upper()
                        
                        # Si cambiamos de firma POSITIVA a CONTRATISTA, agregar más espacio
                        if ("POSITIVA" in texto_actual or "VICEPRESIDENTE" in texto_actual) and \
                           ("CONTRATISTA" in texto_siguiente or "XXXXX" in texto_siguiente):
                            documento_final.add_paragraph()
                            documento_final.add_paragraph()  # Espacio extra entre las dos firmas
                            documento_final.add_paragraph()  # Espacio extra entre las dos firmas

            # Actualizar progreso
            imprimir_dinamico(f"{progreso} - Guardando archivos...")

            # Guardar el documento final
            documento_final.save(ruta_archivo_word)
            documentos_generados += 1

            # Copiar archivo Word a OneDrive
            copiar_word_a_onedrive(ruta_archivo_word)
            archivos_copiados_onedrive += 1
            
            # Registrar éxito en log
            agregar_log(nit, "Procesado correctamente")
            
            # Mostrar resultado final para este NIT
            limpiar_linea()
            nombre_corto = nombre_proveedor[:30] + "..." if len(nombre_proveedor) > 30 else nombre_proveedor
            print(f"[{index + 1}/{total_nits}] ✅ {nit} - {nombre_corto} ({indicadores_procesados} indicadores)")
            
        except Exception as e:
            # Registrar error en log
            error_msg = str(e)
            if "NIT no encontrado" in error_msg:
                agregar_log(nit, "Error - NIT no encontrado en proveedores.xlsx")
            elif "base de datos" in error_msg:
                agregar_log(nit, "Error - Conexión base de datos")
            elif "Categoría" in error_msg and "no encontrada" in error_msg:
                agregar_log(nit, f"Error - Categoría '{categoria_ccm_borrador}' no encontrada")
            elif "tabla plantilla" in error_msg:
                agregar_log(nit, "Error - No se encontró tabla plantilla para fichas")
            elif "imagen" in error_msg:
                agregar_log(nit, "Error - Problema con imagen f.png")
            elif "OneDrive" in error_msg:
                agregar_log(nit, "Error - Problema copiando a OneDrive")
            else:
                agregar_log(nit, f"Error - {error_msg[:50]}...")
            
            # Mostrar error en consola
            limpiar_linea()
            print(f"[{index + 1}/{total_nits}] ❌ {nit} - {error_msg[:60]}...")

    # Guardar archivo de logs
    limpiar_linea()
    print(f"\n{'='*60}")
    print(f"📊 PROCESO FINALIZADO")
    print(f"{'='*60}")
    
    # Guardar logs y mostrar información
    ruta_log = guardar_logs()
    if ruta_log:
        print(f"📋 Log guardado: {os.path.basename(ruta_log)}")
    
    print(f"✅ Documentos Word generados: {documentos_generados}")
    print(f"📁 Archivos copiados a OneDrive: {archivos_copiados_onedrive}")
    print(f"📝 Tasa de éxito: {documentos_generados}/{total_nits} ({(documentos_generados/total_nits*100):.1f}%)" if total_nits > 0 else "")
    
    # Mostrar resumen de errores desde los logs
    errores_por_tipo = {}
    for log in logs_procesamiento:
        if "Error" in log['estado']:
            tipo_error = log['estado'].split(' - ')[1] if ' - ' in log['estado'] else log['estado']
            errores_por_tipo[tipo_error] = errores_por_tipo.get(tipo_error, 0) + 1
    
    if errores_por_tipo:
        print(f"\n📋 Resumen de errores:")
        for tipo_error, cantidad in errores_por_tipo.items():
            print(f"   • {tipo_error}: {cantidad} casos")
    
    if documentos_generados > 0:
        print(f"\n💡 Para generar PDFs, ejecuta el script convertidor_pdf.py")
        print(f"📂 Archivos guardados en: {BASE_DIR_OUTPUT}")

# --- Ejecución del script ---
if __name__ == "__main__":
    try:
        generar_documentos_indicadores()
    except KeyboardInterrupt:
        limpiar_linea()
        print("\n❌ Proceso interrumpido por el usuario")
        # Guardar logs parciales
        if logs_procesamiento:
            ruta_log = guardar_logs()
            if ruta_log:
                print(f"📋 Log parcial guardado: {os.path.basename(ruta_log)}")
    except Exception as e:
        limpiar_linea()
        print(f"\n❌ Error inesperado: {e}")
        # Guardar logs parciales
        if logs_procesamiento:
            ruta_log = guardar_logs()
            if ruta_log:
                print(f"📋 Log parcial guardado: {os.path.basename(ruta_log)}")
        import traceback
        traceback.print_exc()



### Funcional completo con pdf y word

# import pandas as pd
# from docx import Document
# import mysql.connector
# from datetime import date, datetime
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.shared import Inches, Cm
# from docx.oxml import OxmlElement
# from docx.oxml.ns import qn
# from docx.enum.table import WD_TABLE_ALIGNMENT
# from docx.shared import Pt
# import shutil
# import os

# # --- Configuración de rutas ---
# BASE_DIR = r'D:\Proyectos\Indicadores\Resources'
# PATH_LOGS = r'D:\Proyectos\Indicadores\Logs'
# BASE_DIR_OUTPUT = r'D:\Proyectos\Indicadores\Salidas'
# RUTA_LOGS = os.path.join(PATH_LOGS)
# RUTA_IMAGEN = os.path.join(BASE_DIR, 'f.png')
# RUTA_READING = os.path.join(BASE_DIR, 'reading.xlsx')
# RUTA_PROVEEDORES = os.path.join(BASE_DIR, 'proveedores.xlsx')
# # RUTA_PROVEEDORES = os.path.join(BASE_DIR, 'nuevos.xlsx')
# RUTA_INDICADORES = os.path.join(BASE_DIR, 'indicadores_anexo_9.xlsx')
# RUTA_BORRADOR = os.path.join(BASE_DIR, 'borrador_indicadores.docx')
# # --- Ruta One drive para almacenar copia de los Anexos 9 generados ---
# COPIE_DIR = os.path.join(r'C:\Users\JORGEEDILSONVEGAACOS\One Drive Analista Red 11\OneDrive - Positiva Compañia de Seguros S. A\Indicadores\IND New')

# # --- Configuración de la Base de Datos ---
# db_config = {
#     'host': 'localhost',
#     'port': 3306,
#     'user': 'root',
#     'password': '!1q2w3e4r5t6y7u8i9o0p',
#     'database': 'data_tarifas_reps'
# }

# def formatear_fecha(valor, fmt="%Y-%m-%d"):
#     """
#     Devuelve la fecha en formato yyyy/mm/dd.
#     Acepta pandas.Timestamp, datetime/date, string o seriales de Excel.
#     Si no puede convertir, devuelve str(valor).
#     """
#     if valor is None or (isinstance(valor, float) and pd.isna(valor)):
#         return ""
#     if isinstance(valor, (pd.Timestamp, datetime, date)):
#         return pd.to_datetime(valor).strftime(fmt)
#     dt = pd.to_datetime(valor, errors="coerce")
#     if pd.isna(dt):
#         return str(valor) if pd.notna(valor) else ""
#     return dt.strftime(fmt)

# # --- Funciones Auxiliares ---
# def obtener_codigos_reps_habilitados(nit):
#     """
#     Consulta la base de datos para obtener los códigos REPS habilitados para un NIT.
#     """
#     codigos_habilitados = set()
#     try:
#         conn = mysql.connector.connect(**db_config)
#         cursor = conn.cursor()
#         query = f"SELECT DISTINCT cr.serv_codigo FROM consolidado_reps cr WHERE cr.nit = '{nit}'"
#         cursor.execute(query)
#         for (serv_codigo,) in cursor:
#             codigos_habilitados.add(str(serv_codigo).strip())
#         cursor.close()
#         conn.close()
#         print(f"Códigos REPS habilitados para NIT {nit}: {len(codigos_habilitados)} códigos encontrados")
#     except mysql.connector.Error as err:
#         print(f"Error al conectar a la base de datos: {err}")
#     return codigos_habilitados

# def es_numerico(s):
#     """
#     Verifica si una cadena de texto es completamente numérica.
#     """
#     return str(s).isdigit()

# def consolidar_indicadores_duplicados(indicadores_aplicables_df, codigos_reps_habilitados):
#     """
#     Consolida indicadores que tienen el mismo nombre_indicador concatenando cod_propio y cod_reps.
#     Retorna una lista de diccionarios con los indicadores consolidados.
#     """
#     indicadores_consolidados = {}
    
#     for ind_index, ind_row in indicadores_aplicables_df.iterrows():
#         cod_reps = str(ind_row['cod_reps']).strip()
        
#         # Validar si el indicador debe aplicarse
#         aplicar_indicador = False
#         if es_numerico(cod_reps):
#             if cod_reps in codigos_reps_habilitados:
#                 aplicar_indicador = True
#         else:
#             aplicar_indicador = True
        
#         if aplicar_indicador:
#             nombre_indicador = str(ind_row['nombre_indicador']).strip()
#             cod_propio = str(ind_row['cod_propio']).strip()
            
#             # Si ya existe un indicador con este nombre
#             if nombre_indicador in indicadores_consolidados:
#                 # Concatenar cod_propio
#                 cod_propio_existente = indicadores_consolidados[nombre_indicador]['cod_propio']
#                 if cod_propio not in cod_propio_existente:
#                     indicadores_consolidados[nombre_indicador]['cod_propio'] = f"{cod_propio_existente}, {cod_propio}"
                
#                 # Concatenar cod_reps
#                 cod_reps_existente = indicadores_consolidados[nombre_indicador]['cod_reps']
#                 if cod_reps not in cod_reps_existente:
#                     indicadores_consolidados[nombre_indicador]['cod_reps'] = f"{cod_reps_existente}, {cod_reps}"
#             else:
#                 # Crear nuevo indicador consolidado
#                 indicador_consolidado = dict(ind_row)
#                 indicador_consolidado['cod_propio'] = cod_propio
#                 indicador_consolidado['cod_reps'] = cod_reps
#                 indicadores_consolidados[nombre_indicador] = indicador_consolidado
    
#     # Convertir el diccionario a lista para mantener compatibilidad con el resto del código
#     return list(indicadores_consolidados.values())

# def agregar_imagen_ficha_tecnica(documento):
#     """
#     Agrega la imagen f.png como encabezado de ficha técnica.
#     """
#     try:
#         if os.path.exists(RUTA_IMAGEN):
#             # Crear un nuevo párrafo para la imagen
#             paragraph = documento.add_paragraph()
#             paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
#             # Eliminar espaciado antes y después del párrafo
#             paragraph_format = paragraph.paragraph_format
#             paragraph_format.space_before = Pt(0)
#             paragraph_format.space_after = Pt(0)
            
#             # Agregar la imagen al párrafo
#             run = paragraph.add_run()
#             # Ancho igual al de la tabla (4.8 + 11.2 = 16 cm) y alto de 0.7 cm
#             run.add_picture(RUTA_IMAGEN, width=Cm(16), height=Cm(0.7))
            
#             return True
#         else:
#             print(f"Advertencia: No se encontró el archivo f.png en la ruta: {RUTA_IMAGEN}")
#             return False
            
#     except Exception as e:
#         print(f"Error al agregar imagen f.png: {e}")
#         return False

# def agregar_bordes_tabla(tabla):
#     """
#     Agrega solo bordes horizontales internos a la tabla (entre filas).
#     """
#     try:
#         # Obtener el elemento XML de la tabla
#         tbl = tabla._element
        
#         # Crear elemento de bordes
#         tblPr = tbl.find(qn('w:tblPr'))
#         if tblPr is None:
#             tblPr = OxmlElement('w:tblPr')
#             tbl.insert(0, tblPr)
        
#         # Crear bordes de tabla
#         tblBorders = OxmlElement('w:tblBorders')
        
#         # Definir estilos de borde
#         border_style = "single"
#         border_size = "4"  # Tamaño del borde
#         border_color = "000000"  # Color negro
        
#         # Solo crear borde horizontal interno (entre filas)
#         insideH = OxmlElement('w:insideH')
#         insideH.set(qn('w:val'), border_style)
#         insideH.set(qn('w:sz'), border_size)
#         insideH.set(qn('w:color'), border_color)
#         tblBorders.append(insideH)
        
#         # Eliminar bordes existentes si los hay
#         existing_borders = tblPr.find(qn('w:tblBorders'))
#         if existing_borders is not None:
#             tblPr.remove(existing_borders)
            
#         tblPr.append(tblBorders)
        
#     except Exception as e:
#         print(f"Advertencia: No se pudieron agregar bordes a la tabla: {e}")

# def copiar_formato_tabla_simple(tabla_origen, tabla_destino):
#     """
#     Copia el formato básico de una tabla origen a una tabla destino con solo 2 columnas.
#     """
#     try:
#         # Copiar el estilo de la tabla
#         if tabla_origen.style:
#             tabla_destino.style = tabla_origen.style
        
#         # Configurar el ancho total de la tabla
#         tabla_destino.autofit = False
        
#         # Ajustar el ancho de las columnas: 30% para títulos, 70% para contenido
#         if len(tabla_destino.columns) == 2:
#             try:
#                 from docx.shared import Cm
#                 # Usar centímetros para mayor precisión
#                 # Ancho total aproximado de 16cm, 30% = 4.8cm, 70% = 11.2cm
#                 tabla_destino.columns[0].width = Cm(4.8)   # 30%
#                 tabla_destino.columns[1].width = Cm(11.2)  # 70%
                
#                 # Forzar el ancho a nivel XML para mayor efectividad
#                 tbl = tabla_destino._element
#                 tblGrid = tbl.find(qn('w:tblGrid'))
#                 if tblGrid is not None:
#                     gridCols = tblGrid.findall(qn('w:gridCol'))
#                     if len(gridCols) >= 2:
#                         gridCols[0].set(qn('w:w'), '2304')  # 30% en twips (4.8cm * 480)
#                         gridCols[1].set(qn('w:w'), '5376')  # 70% en twips (11.2cm * 480)
                        
#             except Exception as e:
#                 print(f"No se pudo ajustar el ancho de las columnas: {e}")
                
#     except Exception as e:
#         print(f"Advertencia: No se pudo copiar completamente el formato de la tabla: {e}")

# def copiar_formato_tabla(tabla_origen, tabla_destino):
#     """
#     Copia el formato de una tabla origen a una tabla destino.
#     """
#     try:
#         # Copiar el estilo de la tabla
#         if tabla_origen.style:
#             tabla_destino.style = tabla_origen.style
        
#         # Copiar propiedades de celdas
#         for r_idx, row_origen in enumerate(tabla_origen.rows):
#             if r_idx < len(tabla_destino.rows):
#                 for c_idx, cell_origen in enumerate(row_origen.cells):
#                     if c_idx < len(tabla_destino.rows[r_idx].cells):
#                         cell_destino = tabla_destino.rows[r_idx].cells[c_idx]
#                         # Copiar el ancho de la columna si es posible
#                         try:
#                             cell_destino.width = cell_origen.width
#                         except:
#                             pass
#     except Exception as e:
#         print(f"Advertencia: No se pudo copiar completamente el formato de la tabla: {e}")

# def reemplazar_placeholders_en_tabla(tabla, indicador, supervisor, ccm):
#     """
#     Reemplaza los placeholders en una tabla con los datos del indicador.
#     """
#     # Diccionario de reemplazos
#     reemplazos = {
#         '[tipo_proveedor]': str(ccm),
#         '[tipo_indicador]': str(indicador.get('tipo_indicador', '')),
#         '[categoria]': str(indicador.get('categoria', '')),
#         '[nombre_indicador]': str(indicador.get('nombre_indicador', '')),
#         '[fecha_creacion_indicador]': str(indicador.get('fecha_creacion_indicador', '')),
#         '[cod_res_256_de_2016]': str(indicador.get('cod_res_256_de_2016', '')),
#         '[cod_propio]': str(indicador.get('cod_propio', '')),
#         '[descripcion]': str(indicador.get('descripcion', '')),
#         '[formula]': str(indicador.get('formula', '')),
#         '[numerador]': str(indicador.get('numerador', '')),
#         '[fuente_numerador]': str(indicador.get('fuente_numerador', '')),
#         '[denominador]': str(indicador.get('denominador', '')),
#         '[fuente_denominador]': str(indicador.get('fuente_denominador', '')),
#         '[unidad_de_medida]': str(indicador.get('unidad_de_medida', '')),
#         '[meta]': str(indicador.get('meta', '')),
#         '[periodicidad]': str(indicador.get('periodicidad', '')),
#         '[progresividad]': str(indicador.get('progresividad', '')),
#         '[observaciones]': str(indicador.get('observaciones', '')),
#         '[metodologia]': str(indicador.get('metodologia', '')),
#         '[exclusiones]': str(indicador.get('exclusiones', '')),
#         '[responsable]': str(indicador.get('responsable', '')),
#         '[cod_reps]': str(indicador.get('cod_reps', '')),
#         '[grupo_indicador]': str(indicador.get('grupo_indicador', ''))
#     }
    
#     # Reemplazar en todas las celdas
#     for row in tabla.rows:
#         for cell in row.cells:
#             for paragraph in cell.paragraphs:
#                 texto_original = paragraph.text
#                 for placeholder, valor in reemplazos.items():
#                     texto_original = texto_original.replace(placeholder, valor)
#                 paragraph.text = texto_original

# def llenar_tabla_proveedor(tabla, proveedor_data):
#     """
#     Llena la tabla de información del proveedor.
#     """
#     # --- FECHAS del proveedor formateadas ---
#     fin_vig = proveedor_data.get('fin_vigencia', '')

#     # Mapeo de placeholders a valores del proveedor
#     reemplazos_proveedor = {
#         '[sucursal]': str(proveedor_data.get('sucursal', '')),
#         '[nombre]': str(proveedor_data.get('nombre', '')),
#         '[nit]': str(proveedor_data.get('nit', '')),
#         '[numero_contrato]': str(proveedor_data.get('numero_contrato', '')),
#         '[year_contrato]': str(proveedor_data.get('year_contrato', '')),
#         '[fin_vigencia]': formatear_fecha(fin_vig),
#         '[categoría_cuentas_medicas]': str(proveedor_data.get('categoria_cuentas_medicas', '')),
#         '[supervisor]': 'Supervisor y/o interventoría designada'
#     }
    
#     # Reemplazar en todas las celdas de la tabla
#     for row in tabla.rows:
#         for cell in row.cells:
#             for paragraph in cell.paragraphs:
#                 texto_original = paragraph.text
#                 for placeholder, valor in reemplazos_proveedor.items():
#                     texto_original = texto_original.replace(placeholder, valor)
#                 paragraph.text = texto_original

# def agregar_salto_de_pagina(documento):
#     """
#     Agrega un salto de página al documento.
#     """
#     paragraph = documento.add_paragraph()
#     run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
#     fldChar = OxmlElement('w:br')
#     fldChar.set(qn('w:type'), 'page')
#     run._r.append(fldChar)

# def extraer_contenido_cierre_y_firmas(documento_borrador):
#     """
#     Extrae el párrafo de cierre y las firmas del documento borrador.
#     Retorna una tupla: (parrafo_cierre, lista_de_paragrafos_firmas)
#     """
#     parrafo_cierre = None
#     parrafos_firmas = []
#     encontro_cierre = False
    
#     for paragraph in documento_borrador.paragraphs:
#         texto = paragraph.text.strip()
        
#         # Buscar el párrafo de cierre
#         if "Positiva compañía de seguros, mediante la firma de interventoría" in texto:
#             parrafo_cierre = paragraph  # Guardar el párrafo completo, no solo el texto
#             encontro_cierre = True
#             continue
        
#         # Si ya encontró el cierre, los siguientes párrafos son las firmas
#         if encontro_cierre and texto:
#             # Guardar el párrafo completo para mantener el formato
#             parrafos_firmas.append(paragraph)
    
#     return parrafo_cierre, parrafos_firmas

# def copiar_parrafo_con_formato(documento_destino, parrafo_origen):
#     """
#     Copia un párrafo completo con todo su formato al documento destino.
#     """
#     try:
#         # Crear nuevo párrafo en el documento destino
#         nuevo_parrafo = documento_destino.add_paragraph()
        
#         # Copiar formato del párrafo
#         nuevo_parrafo.alignment = parrafo_origen.alignment
#         nuevo_parrafo.paragraph_format.space_before = parrafo_origen.paragraph_format.space_before
#         nuevo_parrafo.paragraph_format.space_after = parrafo_origen.paragraph_format.space_after
#         nuevo_parrafo.paragraph_format.line_spacing = parrafo_origen.paragraph_format.line_spacing
#         nuevo_parrafo.paragraph_format.first_line_indent = parrafo_origen.paragraph_format.first_line_indent
#         nuevo_parrafo.paragraph_format.left_indent = parrafo_origen.paragraph_format.left_indent
#         nuevo_parrafo.paragraph_format.right_indent = parrafo_origen.paragraph_format.right_indent
        
#         # Copiar estilo del párrafo si existe
#         if parrafo_origen.style:
#             try:
#                 nuevo_parrafo.style = parrafo_origen.style
#             except:
#                 pass
        
#         # Copiar cada run (fragmento de texto con formato) del párrafo original
#         for run_origen in parrafo_origen.runs:
#             nuevo_run = nuevo_parrafo.add_run(run_origen.text)
            
#             # Copiar formato del run
#             nuevo_run.bold = run_origen.bold
#             nuevo_run.italic = run_origen.italic
#             nuevo_run.underline = run_origen.underline
            
#             # Copiar fuente si existe
#             if run_origen.font.name:
#                 nuevo_run.font.name = run_origen.font.name
#             if run_origen.font.size:
#                 nuevo_run.font.size = run_origen.font.size
#             if run_origen.font.color.rgb:
#                 nuevo_run.font.color.rgb = run_origen.font.color.rgb
        
#         return nuevo_parrafo
        
#     except Exception as e:
#         print(f"Error copiando formato del párrafo: {e}")
#         # Si hay error, al menos copiar el texto
#         return documento_destino.add_paragraph(parrafo_origen.text)

# def verificar_archivos_requeridos():
#     """
#     Verifica que todos los archivos requeridos existan.
#     """
#     archivos_requeridos = {
#         'reading.xlsx': RUTA_READING,
#         'proveedores.xlsx': RUTA_PROVEEDORES, 
#         'indicadores_anexo_9.xlsx': RUTA_INDICADORES,
#         'borrador_indicadores.docx': RUTA_BORRADOR,
#         'f.png': RUTA_IMAGEN
#     }
    
#     archivos_faltantes = []
#     for nombre, ruta in archivos_requeridos.items():
#         if not os.path.exists(ruta):
#             archivos_faltantes.append(f"{nombre} -> {ruta}")
    
#     if archivos_faltantes:
#         print("Error: No se encontraron los siguientes archivos:")
#         for archivo in archivos_faltantes:
#             print(f"  - {archivo}")
#         return False
    
#     print("Todos los archivos requeridos están disponibles.")
#     return True

# def copiar_word_a_onedrive(ruta_word):
#     """
#     Copia el archivo Word a la carpeta de OneDrive.
#     """
#     try:
#         os.makedirs(COPIE_DIR, exist_ok=True)
#         destino = os.path.join(COPIE_DIR, os.path.basename(ruta_word))
#         shutil.copy2(ruta_word, destino)
#         print(f"  📁 Word copiado a OneDrive: {os.path.basename(ruta_word)}")
#         return True
#     except Exception as e:
#         print(f"  ❌ Error copiando Word a OneDrive: {e}")
#         return False

# # --- Función Principal ---
# def generar_documentos_indicadores():
#     """
#     Función principal que genera los documentos con indicadores.
#     """
#     print("🚀 GENERADOR DE DOCUMENTOS DE INDICADORES (Solo Word)")
#     print("=" * 60)
#     print(f"📁 Directorio base: {BASE_DIR}")
#     print(f"📁 Directorio salida: {BASE_DIR_OUTPUT}")
#     print(f"📁 OneDrive: {COPIE_DIR}")
    
#     # 1. Verificar archivos necesarios
#     if not verificar_archivos_requeridos():
#         return
    
#     # 2. Crear directorios de salida si no existen
#     os.makedirs(BASE_DIR_OUTPUT, exist_ok=True)
#     os.makedirs(COPIE_DIR, exist_ok=True)
    
#     # 3. Cargar los archivos de Excel
#     try:
#         print("\n🔄 Cargando archivos de Excel...")
#         df_borrador = pd.read_excel(RUTA_READING, sheet_name='quering')
#         df_proveedores = pd.read_excel(RUTA_PROVEEDORES, sheet_name='proveedores')
#         df_indicadores_anexo_9 = pd.ExcelFile(RUTA_INDICADORES)
#         df_hom_ccm_ind = df_indicadores_anexo_9.parse('hom_ccm_ind')
#         df_indicadores = df_indicadores_anexo_9.parse('indicadores')
#         print("✅ Archivos cargados exitosamente")
#     except Exception as e:
#         print(f"❌ Error al cargar archivos: {e}")
#         return

#     # 4. Cargar el documento borrador y extraer contenido de cierre y firmas
#     try:
#         doc_borrador_plantilla = Document(RUTA_BORRADOR)
#         parrafo_cierre, parrafos_firmas = extraer_contenido_cierre_y_firmas(doc_borrador_plantilla)
#         print("✅ Documento borrador cargado exitosamente")
#         print(f"✅ Extraído párrafo de cierre y {len(parrafos_firmas)} párrafos de firmas")
#     except Exception as e:
#         print(f"❌ Error al cargar el documento borrador: {e}")
#         return

#     # 5. Iterar sobre cada NIT del borrador
#     documentos_generados = 0
#     archivos_copiados_onedrive = 0
    
#     total_nits = len(df_borrador)
#     print(f"\n🔄 Procesando {total_nits} proveedores...")
    
#     for index, row_borrador in df_borrador.iterrows():
#         nit = str(row_borrador['nit']).strip()
#         categoria_ccm_borrador = str(row_borrador['categoria_cuentas_medicas']).strip()
        
#         print(f"\n[{index + 1}/{total_nits}] 🏢 Procesando NIT: {nit}")

#         # Buscar el NIT en el archivo de proveedores
#         proveedor_data = df_proveedores[df_proveedores['nit'].astype(str).str.strip() == nit]
#         if proveedor_data.empty:
#             print(f"⚠️  Advertencia: No se encontró el NIT {nit} en proveedores.xlsx. Se omite.")
#             continue

#         proveedor_row = proveedor_data.iloc[0]
#         supervisor_proveedor = str(proveedor_row['supervisor'])
#         ccm = str(proveedor_row['categoria_cuentas_medicas'])

#         # Generar el nombre de archivo de salida
#         nombre_proveedor = str(proveedor_row['nombre'])
#         num_contrato = str(proveedor_row['numero_contrato'])
#         year_contrato = str(proveedor_row['year_contrato'])
#         nit_proveedor = str(proveedor_row['nit'])
#         nombre_archivo_word = f"Anexo 9 {num_contrato}-{year_contrato} {nombre_proveedor} {nit_proveedor}.docx"
        
#         # Crear la ruta completa del archivo de salida
#         ruta_archivo_word = os.path.join(BASE_DIR_OUTPUT, nombre_archivo_word)

#         try:
#             # Crear una copia del borrador original
#             shutil.copyfile(RUTA_BORRADOR, ruta_archivo_word)
#             documento_final = Document(ruta_archivo_word)

#             # Llenar la información del proveedor (primera tabla)
#             if len(documento_final.tables) > 0:
#                 tabla_info_proveedor = documento_final.tables[0]
#                 llenar_tabla_proveedor(tabla_info_proveedor, proveedor_row)

#             # Limpiar párrafos de cierre y firmas del documento original
#             # Necesitamos eliminar desde el párrafo de cierre hacia abajo
#             paragrafos_a_eliminar = []
#             eliminar_resto = False
            
#             for paragraph in documento_final.paragraphs:
#                 texto = paragraph.text.strip()
                
#                 # Si encuentra el párrafo de cierre, marcar para eliminar este y todos los siguientes
#                 if "Positiva compañía de seguros, mediante la firma de interventoría" in texto:
#                     eliminar_resto = True
                
#                 # Si ya marcamos para eliminar, agregar todos los párrafos siguientes
#                 if eliminar_resto:
#                     paragrafos_a_eliminar.append(paragraph)

#             # Eliminar los párrafos encontrados
#             for paragraph in paragrafos_a_eliminar:
#                 p = paragraph._element
#                 p.getparent().remove(p)

#             # Encontrar y validar los indicadores
#             grupo_ind_data = df_hom_ccm_ind[df_hom_ccm_ind['categoria_cuentas_medicas'].astype(str).str.strip() == categoria_ccm_borrador]
#             if grupo_ind_data.empty:
#                 print(f"⚠️  Advertencia: No se encontró el grupo_indicador para la categoría '{categoria_ccm_borrador}'. Se omite.")
#                 continue

#             grupo_indicador = str(grupo_ind_data.iloc[0]['grupo_indicador']).strip()
            
#             # Definir grupos a buscar
#             grupos_a_buscar = [grupo_indicador, 'TODOS LOS PROVEEDORES']
#             if grupo_indicador == 'IPS':
#                 grupos_a_buscar.extend(['IPS CON RHB', 'IPS-RHI-RIESGO BIOLOGICO'])
#             elif grupo_indicador == 'ALIADO RHI':
#                 grupos_a_buscar.extend(['IPS-RHI-RIESGO BIOLOGICO'])
#             elif grupo_indicador == 'RIESGO BIOLOGICO':
#                 grupos_a_buscar.extend(['IPS-RHI-RIESGO BIOLOGICO'])
#             elif grupo_indicador == 'MEDICAMENTOS':
#                 grupos_a_buscar.extend(['MEDICAMENTOS'])

#             indicadores_aplicables_df = df_indicadores[df_indicadores['grupo_indicador'].astype(str).str.strip().isin(grupos_a_buscar)]
#             codigos_reps_habilitados = obtener_codigos_reps_habilitados(nit)

#             # NUEVA FUNCIONALIDAD: Consolidar indicadores con mismo nombre_indicador
#             indicadores_consolidados = consolidar_indicadores_duplicados(indicadores_aplicables_df, codigos_reps_habilitados)
            
#             # Obtener la tabla plantilla para las fichas (la última tabla del documento)
#             if len(documento_final.tables) > 1:
#                 ficha_plantilla = documento_final.tables[-1]
                
#                 # Remover la tabla plantilla original para evitar duplicados
#                 tabla_elemento = ficha_plantilla._element
#                 tabla_elemento.getparent().remove(tabla_elemento)
#             else:
#                 print("❌ Error: No se encontró la tabla plantilla para las fichas")
#                 continue

#             # Generar fichas de indicadores consolidados
#             indicadores_procesados = 0
            
#             for indicador in indicadores_consolidados:
#                 nombre_indicador = str(indicador['nombre_indicador']).strip()
#                 cod_propio = str(indicador['cod_propio']).strip()
#                 cod_reps = str(indicador['cod_reps']).strip()
                
#                 print(f"  📝 Creando ficha consolidada para: {nombre_indicador}")
#                 if ',' in cod_propio:
#                     print(f"      🔗 Códigos propios consolidados: {cod_propio}")
#                 if ',' in cod_reps:
#                     print(f"      🔗 Códigos REPS consolidados: {cod_reps}")
                
#                 # Agregar la imagen de "ficha técnica" antes de cada tabla
#                 agregar_imagen_ficha_tecnica(documento_final)
                
#                 # Crear nueva tabla con solo 2 columnas (etiqueta y valor)
#                 # Sin espacio entre la imagen y la tabla
#                 nueva_tabla = documento_final.add_table(
#                     rows=len(ficha_plantilla.rows),
#                     cols=2  # Solo 2 columnas: etiqueta y valor
#                 )
                
#                 # Copiar formato de la tabla original
#                 copiar_formato_tabla_simple(ficha_plantilla, nueva_tabla)
                
#                 # Agregar bordes a la tabla
#                 agregar_bordes_tabla(nueva_tabla)
                
#                 # Copiar solo el contenido de las primeras 2 columnas de la plantilla
#                 # Comenzar desde la fila 0 para evitar filas vacías
#                 filas_copiadas = 0
#                 for r_idx, row_plantilla in enumerate(ficha_plantilla.rows):
#                     # Verificar si la fila tiene contenido útil (no está vacía)
#                     tiene_contenido = False
#                     for c_idx in range(min(2, len(row_plantilla.cells))):
#                         if row_plantilla.cells[c_idx].text.strip():
#                             tiene_contenido = True
#                             break
                    
#                     # Solo copiar filas que tengan contenido
#                     if tiene_contenido and filas_copiadas < len(nueva_tabla.rows):
#                         for c_idx in range(min(2, len(row_plantilla.cells))):
#                             if c_idx < len(nueva_tabla.rows[filas_copiadas].cells):
#                                 nueva_tabla.cell(filas_copiadas, c_idx).text = row_plantilla.cells[c_idx].text.strip()
#                         filas_copiadas += 1
                
#                 # Si sobran filas en la nueva tabla, eliminarlas
#                 while len(nueva_tabla.rows) > filas_copiadas:
#                     row_to_remove = nueva_tabla.rows[-1]._element
#                     row_to_remove.getparent().remove(row_to_remove)

#                 # Reemplazar placeholders con datos del indicador consolidado
#                 reemplazar_placeholders_en_tabla(nueva_tabla, indicador, supervisor_proveedor, ccm)
                
#                 # Agregar espacio entre fichas
#                 documento_final.add_paragraph()
#                 # documento_final.add_paragraph()  # Doble espacio para mejor separación
                
#                 indicadores_procesados += 1

#             # Agregar el párrafo de cierre y las firmas al final del documento
#             if parrafo_cierre:
#                 # Agregar espacio antes del párrafo de cierre
#                 # documento_final.add_paragraph()
                
#                 # Agregar el párrafo de cierre con su formato original
#                 copiar_parrafo_con_formato(documento_final, parrafo_cierre)
                
#                 # Agregar espacio antes de las firmas
#                 documento_final.add_paragraph()
#                 # documento_final.add_paragraph()  # Espacio extra para las firmas
                
#                 # Agregar cada párrafo de firma con su formato original completo
#                 for i, parrafo_firma in enumerate(parrafos_firmas):
#                     copiar_parrafo_con_formato(documento_final, parrafo_firma)
                    
#                     # Si no es la última firma, agregar espacio entre firmas
#                     # Detectar cambio de firma (de POSITIVA a CONTRATISTA)
#                     if i < len(parrafos_firmas) - 1:
#                         texto_actual = parrafo_firma.text.strip().upper()
#                         texto_siguiente = parrafos_firmas[i + 1].text.strip().upper()
                        
#                         # Si cambiamos de firma POSITIVA a CONTRATISTA, agregar más espacio
#                         if ("POSITIVA" in texto_actual or "VICEPRESIDENTE" in texto_actual) and \
#                            ("CONTRATISTA" in texto_siguiente or "XXXXX" in texto_siguiente):
#                             documento_final.add_paragraph()
#                             documento_final.add_paragraph()  # Espacio extra entre las dos firmas
#                             documento_final.add_paragraph()  # Espacio extra entre las dos firmas

#             # Guardar el documento final
#             documento_final.save(ruta_archivo_word)
#             print(f"  ✅ Word guardado: {os.path.basename(ruta_archivo_word)}")
#             print(f"  📊 Indicadores procesados: {indicadores_procesados}")
#             documentos_generados += 1

#             # Copiar archivo Word a OneDrive
#             if copiar_word_a_onedrive(ruta_archivo_word):
#                 archivos_copiados_onedrive += 1
            
#         except Exception as e:
#             print(f"❌ Error procesando NIT {nit}: {e}")
#             continue

#     # Mostrar estadísticas finales
#     print(f"\n{'='*60}")
#     print(f"📊 PROCESO FINALIZADO")
#     print(f"{'='*60}")
#     print(f"✅ Documentos Word generados: {documentos_generados}")
#     print(f"📁 Archivos copiados a OneDrive: {archivos_copiados_onedrive}")
#     print(f"📝 Tasa de éxito: {documentos_generados}/{total_nits} ({(documentos_generados/total_nits*100):.1f}%)" if total_nits > 0 else "")
    
#     if documentos_generados > 0:
#         print(f"\n💡 Para generar PDFs, ejecuta el script convertidor_pdf.py")
#         print(f"📂 Archivos guardados en: {BASE_DIR_OUTPUT}")

# # --- Ejecución del script ---
# if __name__ == "__main__":
#     try:
#         generar_documentos_indicadores()
#     except KeyboardInterrupt:
#         print("\n❌ Proceso interrumpido por el usuario")
#     except Exception as e:
#         print(f"\n❌ Error inesperado: {e}")
#         import traceback
#         traceback.print_exc()